"""
Genera un .docx con las plantillas elegidas del proyecto Web CATRACHO.

Uso:
    python scripts/build-plantillas-docx.py

Salida:
    Plantillas_Elegidas_Web_CATRACHO.docx en la raíz del workspace.
"""
from __future__ import annotations
import os
import sys
from pathlib import Path
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.parent
ELEGIDAS = ROOT / "Plantillas" / "Elegidas"
TMP_PNG = ROOT / "PaginaWebCatracho" / ".tmp_plantillas_png"
OUTPUT = ROOT / "Plantillas_Elegidas_Web_CATRACHO.docx"

TMP_PNG.mkdir(exist_ok=True)


def svg_to_png(svg_path: Path, dpi: int = 144) -> Path:
    """Convierte un SVG a PNG usando svglib + reportlab. Retorna la ruta del PNG."""
    out = TMP_PNG / (svg_path.stem + ".png")
    if out.exists() and out.stat().st_mtime >= svg_path.stat().st_mtime:
        return out
    drawing = svg2rlg(str(svg_path))
    scale = dpi / 72.0
    drawing.scale(scale, scale)
    drawing.width *= scale
    drawing.height *= scale
    renderPM.drawToFile(drawing, str(out), fmt="PNG")
    return out


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Plantillas Elegidas")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1E, 0xA7, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Web CATRACHO — Cámara de Transporte de Carga de Honduras")
    sub_run.font.size = Pt(16)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_run = desc.add_run(
        "Compilado de las plantillas visuales seleccionadas durante la fase de diseño "
        "del proyecto, presentadas en sus versiones desktop y mobile, junto con el "
        "sistema de diseño (UIKit) que rige los componentes de toda la página."
    )
    desc_run.font.size = Pt(12)
    desc_run.italic = True

    doc.add_page_break()


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0xA7, 0xFF)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_image_centered(doc: Document, image_path: Path, max_width_in: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(max_width_in))


def resolve_image(section_dir: Path, base_name: str, root_preview: str | None = None) -> tuple[Path | None, str]:
    """
    Encuentra la mejor representación visual disponible:
    1. PNG render real (.capture.png) en la carpeta de la sección
    2. PNG preview en la raíz de Elegidas/
    3. SVG convertido a PNG con svglib (último recurso)
    """
    capture = section_dir / f"{base_name}.capture.png"
    if capture.exists():
        return capture, capture.name

    if root_preview:
        preview = ELEGIDAS / root_preview
        if preview.exists():
            return preview, preview.name

    svg = section_dir / f"{base_name}.svg"
    if svg.exists():
        return svg_to_png(svg), svg.name

    return None, ""


def add_section(
    doc: Document,
    title: str,
    folder: str,
    desktop_base: str,
    mobile_base: str,
    desktop_root_preview: str | None = None,
    mobile_root_preview: str | None = None,
) -> None:
    add_section_heading(doc, title, level=1)
    section_dir = ELEGIDAS / folder

    add_section_heading(doc, "Vista desktop", level=2)
    desktop_path, desktop_label = resolve_image(section_dir, desktop_base, desktop_root_preview)
    if desktop_path:
        add_image_centered(doc, desktop_path, max_width_in=6.5)
        add_caption(doc, desktop_label)
    else:
        doc.add_paragraph(f"(Vista desktop no disponible para {folder})")

    add_section_heading(doc, "Vista mobile", level=2)
    mobile_path, mobile_label = resolve_image(section_dir, mobile_base, mobile_root_preview)
    if mobile_path:
        add_image_centered(doc, mobile_path, max_width_in=3.5)
        add_caption(doc, mobile_label)
    else:
        doc.add_paragraph(f"(Vista mobile no disponible para {folder})")

    doc.add_page_break()


def add_uikit(doc: Document) -> None:
    add_section_heading(doc, "Sistema de diseño (UIKit)", level=1)
    doc.add_paragraph(
        "Define las fundaciones visuales que se aplican consistentemente a lo largo "
        "de todas las secciones: paleta de colores, escalas tipográficas, botones, "
        "modales en sus variantes desktop y mobile, y tablas. Estos componentes son "
        "la base sobre la que se construyen las plantillas de cada sección."
    )

    uikit_dir = ELEGIDAS / "UIKit"
    items = [
        ("Colores", "UI KIT COLORES.png"),
        ("Tipografía", "UI KIT TIPOGRAFIA.png"),
        ("Botones", "UI KIT BOTONES.png"),
        ("Modal — desktop", "UI KIT MODAL DESKTOP.png"),
        ("Modal — mobile", "UI KIT MODAL MOBILE.png"),
        ("Tabla", "UI KIT TABLA.png"),
    ]
    for title, png_name in items:
        add_section_heading(doc, title, level=2)
        png_path = uikit_dir / png_name
        if png_path.exists():
            add_image_centered(doc, png_path, max_width_in=6.5)
            add_caption(doc, png_name)
        else:
            doc.add_paragraph(f"(Archivo no encontrado: {png_name})")
    doc.add_page_break()


def main() -> None:
    if not ELEGIDAS.exists():
        sys.exit(f"No existe el directorio: {ELEGIDAS}")

    doc = Document()
    add_title_page(doc)

    add_uikit(doc)

    sections = [
        # (titulo, carpeta, base desktop, base mobile, preview desktop opcional, preview mobile opcional)
        ("Landing", "landing", "CATRACHO_Landing_Desktop_V1", "CATRACHO_Mobile_V2", "landing_desktop_preview.png", "landing_mobile_preview.png"),
        ("Historia", "historia", "CATRACHO_Historia_Desktop_V2", "CATRACHO_Historia_Mobile", None, None),
        ("Misión y Visión", "mv", "CATRACHO_MyV_Desktop_V2", "CATRACHO_MyV_Mobile_V1", None, None),
        ("Servicios", "servicios", "CATRACHO_Servicios_Desktop_V2", "CATRACHO_Servicios_Mobile_V1", None, None),
        ("Requisitos", "requisitos", "CATRACHO_Requisitos_Desktop_V1", "CATRACHO_Requisitos_Mobile_V1", None, None),
        ("Información", "infor", "CATRACHO_Info_Modal_Desktop_V1", "CATRACHO_Info_Modal_Mobile_V2", None, "info_mobile_preview.png"),
        ("Leyes y Otros", "leyes", "CATRACHO_LeyesOtros_Desktop_V2", "CATRACHO_LeyesOtros_Mobile", "leyes_desktop_preview.png", None),
        ("Distancias", "distan", "CATRACHO_Distancias_Desktop_V1", "CATRACHO_Distancias_Mobile", None, "distancias_mobile_preview.png"),
        ("Contáctenos", "contact", "CATRACHO_Contactos_Desktop_V1", "CATRACHO_Contactos_Mobile", None, None),
    ]
    for title, folder, desktop, mobile, dpreview, mpreview in sections:
        print(f"  -> {title}")
        add_section(doc, title, folder, desktop, mobile, dpreview, mpreview)

    doc.save(str(OUTPUT))
    print(f"\nGenerado: {OUTPUT}")


if __name__ == "__main__":
    main()
